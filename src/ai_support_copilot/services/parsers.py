from pathlib import Path

import httpx
from bs4 import BeautifulSoup

from ai_support_copilot.domain.models import SourceType
from ai_support_copilot.services.text import normalize_text


class ParsedDocument:
    def __init__(self, title: str, text: str, source_type: SourceType, metadata: dict) -> None:
        self.title = title
        self.text = text
        self.source_type = source_type
        self.metadata = metadata


class DocumentParser:
    async def parse_bytes(self, filename: str, payload: bytes) -> ParsedDocument:
        suffix = Path(filename).suffix.lower()
        if suffix in {".md", ".markdown"}:
            return ParsedDocument(
                filename, payload.decode("utf-8", "ignore"), SourceType.markdown, {}
            )
        if suffix == ".txt":
            return ParsedDocument(filename, payload.decode("utf-8", "ignore"), SourceType.text, {})
        if suffix == ".pdf":
            return ParsedDocument(filename, await self._parse_pdf(payload), SourceType.pdf, {})
        if suffix == ".docx":
            return ParsedDocument(filename, await self._parse_docx(payload), SourceType.docx, {})
        return ParsedDocument(filename, payload.decode("utf-8", "ignore"), SourceType.text, {})

    async def parse_url(self, url: str, title: str | None = None) -> ParsedDocument:
        async with httpx.AsyncClient(timeout=20, follow_redirects=True) as client:
            response = await client.get(url)
            response.raise_for_status()
        soup = BeautifulSoup(response.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer"]):
            tag.decompose()
        page_title = title or (soup.title.string if soup.title else url)
        return ParsedDocument(
            page_title,
            normalize_text(soup.get_text(separator=" ")),
            SourceType.url,
            {"url": url},
        )

    async def _parse_pdf(self, payload: bytes) -> str:
        from io import BytesIO

        from pypdf import PdfReader

        reader = PdfReader(BytesIO(payload))
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)

    async def _parse_docx(self, payload: bytes) -> str:
        from io import BytesIO

        import docx

        document = docx.Document(BytesIO(payload))
        return "\n".join(paragraph.text for paragraph in document.paragraphs)
